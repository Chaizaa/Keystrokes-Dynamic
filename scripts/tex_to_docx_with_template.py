from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


SRC_TEX = Path("papers/drafts/conference_101719_6pages_no_author.tex")
TEMPLATE_DOCX = Path("papers/references/conference-template-a4.docx")
OUT_DOCX = Path("papers/outputs/conference_101719_6pages_no_author.docx")


def _clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _extract_template_margins(template_path: Path) -> dict[str, float] | None:
    if not template_path.exists():
        return None
    try:
        with zipfile.ZipFile(template_path, "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", "ignore")
    except Exception:
        return None

    sects = re.findall(r"<w:sectPr[\s\S]*?</w:sectPr>", xml)
    if not sects:
        return None

    last = sects[-1]
    mar = re.search(r"<w:pgMar[^>]*/>", last)
    if not mar:
        return None

    attrs = mar.group(0)

    def _get_pt(name: str, default_pt: float) -> float:
        m = re.search(rf'w:{name}="([0-9.]+)pt"', attrs)
        return float(m.group(1)) if m else default_pt

    return {
        "top": _get_pt("top", 54.0),
        "right": _get_pt("right", 44.65),
        "bottom": _get_pt("bottom", 72.0),
        "left": _get_pt("left", 44.65),
    }


def _load_doc_with_template_fallback(template_path: Path) -> Document:
    try:
        return Document(str(template_path))
    except Exception:
        doc = Document()
        margins = _extract_template_margins(template_path)
        if margins:
            for sec in doc.sections:
                sec.top_margin = Inches(margins["top"] / 72.0)
                sec.right_margin = Inches(margins["right"] / 72.0)
                sec.bottom_margin = Inches(margins["bottom"] / 72.0)
                sec.left_margin = Inches(margins["left"] / 72.0)
        else:
            for sec in doc.sections:
                sec.top_margin = Inches(0.75)
                sec.right_margin = Inches(0.62)
                sec.bottom_margin = Inches(1.0)
                sec.left_margin = Inches(0.62)
        return doc


def _replace_nested(cmd: str, text: str) -> str:
    pattern = re.compile(rf"\\{cmd}\{{([^{{}}]*)\}}")
    while True:
        new_text = pattern.sub(r"\1", text)
        if new_text == text:
            return text
        text = new_text


def _clean_latex_text(text: str, cite_index: dict[str, int]) -> str:
    def cite_repl(match: re.Match[str]) -> str:
        keys = [k.strip() for k in match.group(1).split(",") if k.strip()]
        nums = [str(cite_index[k]) for k in keys if k in cite_index]
        return f"[{','.join(nums)}]" if nums else ""

    text = text.replace("\r", "")
    text = re.sub(r"%.*", "", text)
    text = re.sub(r"\\cite\{([^}]*)\}", cite_repl, text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{[^}]*\}", "", text)

    for cmd in ("texttt", "textit", "emph", "mathrm", "mathbf", "operatorname"):
        text = _replace_nested(cmd, text)

    text = text.replace("\\&", "&")
    text = text.replace("\\%", "%")
    text = text.replace("~", " ")

    # Keep simple math readable in Word text flow.
    text = text.replace("$", "")

    # Remove remaining control words but keep their arguments where possible.
    text = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("\\", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _extract_bibliography(tex: str) -> list[tuple[str, str]]:
    bib_block_match = re.search(
        r"\\begin\{thebibliography\}\{[^}]*\}(.*?)\\end\{thebibliography\}",
        tex,
        re.DOTALL,
    )
    if not bib_block_match:
        return []

    bib_block = bib_block_match.group(1)
    entries: list[tuple[str, str]] = []
    for m in re.finditer(
        r"\\bibitem\{([^}]*)\}\s*(.*?)(?=\\bibitem\{|\Z)",
        bib_block,
        re.DOTALL,
    ):
        key = m.group(1).strip()
        raw = m.group(2).strip()
        entries.append((key, raw))
    return entries


def _extract_sections(tex: str) -> list[tuple[str, str]]:
    content_match = re.search(
        r"\\begin\{IEEEkeywords\}.*?\\end\{IEEEkeywords\}(.*?)(?=\\bibliographystyle)",
        tex,
        re.DOTALL,
    )
    if not content_match:
        return []

    content = content_match.group(1)
    sections: list[tuple[str, str]] = []
    sec_matches = list(re.finditer(r"\\section\{([^}]*)\}", content))
    for i, m in enumerate(sec_matches):
        title = m.group(1).strip()
        start = m.end()
        end = sec_matches[i + 1].start() if i + 1 < len(sec_matches) else len(content)
        block = content[start:end]
        sections.append((title, block))
    return sections


def _add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for r in heading.runs:
        if level == 1:
            r.font.size = Pt(12)
            r.bold = True
        else:
            r.font.size = Pt(11)
            r.bold = True


def _add_paragraph(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(18)
    for r in p.runs:
        r.font.size = Pt(11)


def main() -> None:
    if not SRC_TEX.exists():
        raise FileNotFoundError(f"Source TeX not found: {SRC_TEX}")
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(f"Template DOCX not found: {TEMPLATE_DOCX}")

    tex = SRC_TEX.read_text(encoding="utf-8")

    title_match = re.search(r"\\title\{([^}]*)\}", tex, re.DOTALL)
    title = title_match.group(1).strip() if title_match else "Conference Paper"

    abstract_match = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", tex, re.DOTALL)
    abstract = abstract_match.group(1).strip() if abstract_match else ""

    kw_match = re.search(r"\\begin\{IEEEkeywords\}(.*?)\\end\{IEEEkeywords\}", tex, re.DOTALL)
    keywords = kw_match.group(1).strip() if kw_match else ""

    bib_entries = _extract_bibliography(tex)
    cite_index = {key: i + 1 for i, (key, _) in enumerate(bib_entries)}

    doc = _load_doc_with_template_fallback(TEMPLATE_DOCX)
    _clear_document_body(doc)

    _add_title(doc, _clean_latex_text(title, cite_index))

    _add_heading(doc, "Abstract", level=1)
    _add_paragraph(doc, _clean_latex_text(abstract, cite_index))

    kw_para = doc.add_paragraph()
    kw_run = kw_para.add_run(f"Keywords: {_clean_latex_text(keywords, cite_index)}")
    kw_run.italic = True
    kw_run.font.size = Pt(10)
    kw_para.paragraph_format.space_after = Pt(8)

    for sec_title, sec_block in _extract_sections(tex):
        _add_heading(doc, sec_title, level=1)

        block = re.sub(r"\\begin\{table\}.*?\\end\{table\}", "", sec_block, flags=re.DOTALL)
        block = re.sub(r"\\begin\{equation\}.*?\\end\{equation\}", "", block, flags=re.DOTALL)

        # Split by subsections first.
        parts = re.split(r"\\subsection\{([^}]*)\}", block)
        if len(parts) == 1:
            paras = [p.strip() for p in block.split("\n\n") if p.strip()]
            for p in paras:
                cleaned = _clean_latex_text(p, cite_index)
                if cleaned:
                    _add_paragraph(doc, cleaned)
            continue

        intro_part = parts[0].strip()
        if intro_part:
            for p in [x.strip() for x in intro_part.split("\n\n") if x.strip()]:
                cleaned = _clean_latex_text(p, cite_index)
                if cleaned:
                    _add_paragraph(doc, cleaned)

        for i in range(1, len(parts), 2):
            sub_title = parts[i].strip()
            sub_body = parts[i + 1] if i + 1 < len(parts) else ""
            _add_heading(doc, sub_title, level=2)
            for p in [x.strip() for x in sub_body.split("\n\n") if x.strip()]:
                cleaned = _clean_latex_text(p, cite_index)
                if cleaned:
                    _add_paragraph(doc, cleaned)

    _add_heading(doc, "References", level=1)
    for i, (_key, raw_entry) in enumerate(bib_entries, start=1):
        ref = _clean_latex_text(raw_entry, cite_index)
        if not ref:
            continue
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.size = Pt(10)

    doc.save(str(OUT_DOCX))
    print(f"Generated: {OUT_DOCX}")


if __name__ == "__main__":
    main()