from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from fpdf import FPDF


SOURCE_MD = Path("docs/MANUAL_BOOK_NON_TECHNICAL.md")
OUTPUT_DOCX = Path("docs/MANUAL_BOOK_NON_TECHNICAL.docx")
OUTPUT_PDF = Path("docs/MANUAL_BOOK_NON_TECHNICAL.pdf")


def _iter_md_lines(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8").splitlines()


def _is_table_line(stripped: str) -> bool:
    return stripped.startswith("|") and stripped.endswith("|")


def _split_table_cells(row: str) -> list[str]:
    return [cell.strip() for cell in row.strip().strip("|").split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    for cell in cells:
        compact = cell.replace(" ", "")
        if not compact:
            continue
        if not set(compact) <= {"-", ":"}:
            return False
    return True


def _consume_table(lines: list[str], start: int) -> tuple[Optional[list[list[str]]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines):
        stripped = lines[i].strip()
        if not _is_table_line(stripped):
            break
        table_lines.append(stripped)
        i += 1

    if len(table_lines) < 2:
        return None, start

    rows = [_split_table_cells(row) for row in table_lines]
    if len(rows) >= 2 and _is_separator_row(rows[1]):
        rows = [rows[0]] + rows[2:]

    if not rows:
        return None, start

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]
    return normalized_rows, i


def _write_docx(lines: list[str], output_path: Path) -> None:
    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped == "[[PAGE_BREAK]]":
            doc.add_page_break()
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("")
            i += 1
            continue

        if _is_table_line(stripped):
            table_rows, next_i = _consume_table(lines, i)
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = "Table Grid"
                for r_idx, row in enumerate(table_rows):
                    for c_idx, value in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = value
                        if r_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                doc.add_paragraph("")
                i = next_i
                continue

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            level = min(max(level, 1), 9)
            doc.add_heading(text, level=level)
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        if _is_numbered_item(stripped):
            item_text = stripped.split(".", 1)[1].strip()
            doc.add_paragraph(item_text, style="List Number")
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    doc.save(output_path)


def _is_numbered_item(line: str) -> bool:
    if "." not in line:
        return False
    first, _rest = line.split(".", 1)
    return first.isdigit()


def _normalize_for_pdf(text: str, chunk: int = 60) -> str:
    """Insert soft spaces into very long tokens so PDF line-breaker can wrap them."""
    parts = text.split(" ")
    normalized: list[str] = []
    for token in parts:
        if len(token) <= chunk:
            normalized.append(token)
            continue
        slices = [token[i : i + chunk] for i in range(0, len(token), chunk)]
        normalized.append(" ".join(slices))
    return " ".join(normalized)


def _fit_text_for_cell(pdf: FPDF, text: str, max_width: float) -> str:
    candidate = text
    while candidate and pdf.get_string_width(candidate) > max_width - 2:
        if len(candidate) <= 3:
            return "..."
        candidate = candidate[:-4] + "..."
    return candidate


class ManualPDF(FPDF):
    def header(self) -> None:  # type: ignore[override]
        pass

    def footer(self) -> None:  # type: ignore[override]
        self.set_y(-12)
        self.set_font("Helvetica", size=8)
        self.cell(0, 8, f"Page {self.page_no()}", align="C")


def _write_pdf(lines: list[str], output_path: Path) -> None:
    pdf = ManualPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    content_width = pdf.w - pdf.l_margin - pdf.r_margin

    def write_line(text: str, *, heading: bool = False) -> None:
        safe_text = _normalize_for_pdf(text)
        line_height = 8 if heading else 6
        pdf.multi_cell(content_width, line_height, safe_text)

    def write_table(rows: list[list[str]]) -> None:
        if not rows:
            return
        col_count = len(rows[0])
        if col_count == 0:
            return

        col_width = content_width / col_count
        row_height = 7

        for r_idx, row in enumerate(rows):
            if pdf.get_y() > (pdf.h - pdf.b_margin - row_height):
                pdf.add_page()
            for c_idx, cell in enumerate(row):
                text = _fit_text_for_cell(pdf, _normalize_for_pdf(cell), col_width)
                style = "B" if r_idx == 0 else ""
                pdf.set_font("Helvetica", style=style, size=10)
                pdf.cell(col_width, row_height, text, border=1)
            pdf.ln(row_height)
        pdf.ln(2)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped == "[[PAGE_BREAK]]":
            pdf.add_page()
            i += 1
            continue

        if stripped == "---":
            pdf.ln(2)
            i += 1
            continue

        if _is_table_line(stripped):
            table_rows, next_i = _consume_table(lines, i)
            if table_rows:
                write_table(table_rows)
                i = next_i
                continue

        if not stripped:
            pdf.ln(3)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            size = {1: 18, 2: 15, 3: 13, 4: 12}.get(level, 11)
            pdf.set_font("Helvetica", style="B", size=size)
            write_line(text, heading=True)
            pdf.ln(1)
            i += 1
            continue

        if stripped.startswith("- "):
            pdf.set_font("Helvetica", size=11)
            write_line(f"- {stripped[2:].strip()}")
            i += 1
            continue

        if _is_numbered_item(stripped):
            pdf.set_font("Helvetica", size=11)
            write_line(stripped)
            i += 1
            continue

        pdf.set_font("Helvetica", size=11)
        write_line(stripped)
        i += 1

    pdf.output(str(output_path))


def main() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")

    lines = _iter_md_lines(SOURCE_MD)
    _write_docx(lines, OUTPUT_DOCX)
    _write_pdf(lines, OUTPUT_PDF)

    print(f"Generated DOCX: {OUTPUT_DOCX}")
    print(f"Generated PDF : {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
