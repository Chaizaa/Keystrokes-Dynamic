"""
Transform RESEARCH_PAPER_DRAFT.md to conference template DOCX format.
Follows IEEE/conference standards with LaTeX-style mathematical typesetting.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE


def setup_document_styles(doc: Document) -> None:
    """Configure document styles for IEEE/conference format."""
    
    # Margins - conference standard
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_footer_page_numbers(doc: Document) -> None:
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Styling
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)


def add_title_page_conference(doc: Document, title: str, authors: list[str], affiliation: str) -> None:
    """Add conference-style title page."""
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    title_para.paragraph_format.space_before = Pt(24)
    
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Authors
    authors_para = doc.add_paragraph()
    authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_para.paragraph_format.space_after = Pt(6)
    
    authors_run = authors_para.add_run(", ".join(authors))
    authors_run.font.size = Pt(11)
    authors_run.font.name = 'Times New Roman'
    
    # Affiliation
    affil_para = doc.add_paragraph()
    affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_para.paragraph_format.space_after = Pt(24)
    
    affil_run = affil_para.add_run(affiliation)
    affil_run.font.size = Pt(10)
    affil_run.font.italic = True
    affil_run.font.name = 'Times New Roman'


def parse_latex_equations(text: str) -> str:
    """Parse LaTeX equations and mark them for special handling."""
    # For now, keep LaTeX syntax visible - Word will render math differently
    return text


def create_heading(doc: Document, text: str, level: int = 1) -> None:
    """Create IEEE-style heading."""
    
    if level == 1:
        # Main section heading (e.g., "I. INTRODUCTION")
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.all_caps = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    
    elif level == 2:
        # Subsection heading (e.g., "A. Background")
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = 'Times New Roman'
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
    
    elif level == 3:
        # Sub-subsection
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.italic = True
            run.font.name = 'Times New Roman'
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)


def add_body_paragraph(doc: Document, text: str, indent_first: bool = True) -> None:
    """Add body text paragraph with IEEE formatting."""
    
    if not text or not text.strip():
        return
    
    # Clean markdown
    text = text.strip()
    text = re.sub(r"```[\s\S]*?```", "", text)  # Remove code blocks
    text = re.sub(r"\[ISI:.*?\]", "", text)  # Remove ISI placeholders
    text = re.sub(r"\[(\d+)\]", r"[\1]", text)  # Keep citation format
    
    para = doc.add_paragraph()
    
    # First-line indent (IEEE style)
    if indent_first:
        para.paragraph_format.first_line_indent = Inches(0.25)
    
    # Spacing
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def add_table_from_markdown(doc: Document, table_text: str) -> None:
    """Convert markdown table to DOCX table."""
    
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    
    if len(lines) < 2:
        return
    
    # Parse rows
    rows = []
    for i, line in enumerate(lines):
        if i == 1:  # Skip separator
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Format table
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            
            # Clear and add paragraph
            cell.paragraphs[0].text = cell_data
            
            if row_idx == 0:  # Header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # Space after table


def extract_and_format_sections(markdown_content: str) -> dict[str, str]:
    """Extract paper sections from markdown."""
    
    sections = {}
    
    # Extract title
    title_match = re.search(r"^# (.+?)$", markdown_content, re.MULTILINE)
    sections["title"] = title_match.group(1) if title_match else "Keystroke Dynamics Authentication"
    
    # Extract abstract
    abstract_match = re.search(r"## Abstrak\n(.*?)(?=## Abstract|## I\.|$)", markdown_content, re.DOTALL)
    if abstract_match:
        sections["abstract"] = abstract_match.group(1).strip()
    
    # Extract main sections (I-V)
    for roman, num in [("I", 1), ("II", 2), ("III", 3), ("IV", 4), ("V", 5)]:
        pattern = rf"## {roman}\.\s+(.+?)\n(.*?)(?=## [I]|$)"
        match = re.search(pattern, markdown_content, re.DOTALL)
        if match:
            section_title = match.group(1)
            section_content = match.group(2)
            sections[f"section_{num}"] = (section_title, section_content)
    
    return sections


def create_conference_docx(markdown_path: str, output_path: str) -> None:
    """Create conference-format DOCX from markdown paper."""
    
    # Read markdown
    with open(markdown_path, encoding="utf-8") as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    add_footer_page_numbers(doc)
    
    # Extract sections
    sections = extract_and_format_sections(content)
    
    # Title page
    add_title_page_conference(
        doc,
        sections.get("title", "Research Paper"),
        ["Author One", "Author Two"],
        "University Name"
    )
    
    # Abstract (English)
    abstract_heading = doc.add_heading("Abstract", level=1)
    for run in abstract_heading.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
    
    abstract = sections.get("abstract", "")
    if abstract:
        abstract_text = re.sub(r"## Abstract\n", "", abstract)
        abstract_text = abstract_text.split("**Kata Kunci:**")[0].strip()
        add_body_paragraph(doc, abstract_text, indent_first=False)
    
    # Add page break
    doc.add_page_break()
    
    # Main sections
    section_order = [
        ("section_1", "I. INTRODUCTION"),
        ("section_2", "II. RELATED WORK AND METHODOLOGY"),
        ("section_3", "III. DATASET AND EXPERIMENT SETUP"),
        ("section_4", "IV. RESULTS AND DISCUSSION"),
        ("section_5", "V. CONCLUSION"),
    ]
    
    for section_key, default_title in section_order:
        if section_key in sections:
            title, content = sections[section_key]
            
            # Add main heading
            create_heading(doc, f"{default_title.split('. ')[0]}. {title.upper()}", level=1)
            
            # Parse content for tables and paragraphs
            blocks = re.split(r"\n\s*\n+", content)
            
            for block in blocks:
                block = block.strip()
                if not block:
                    continue
                
                # Check if markdown table
                if block.count("|") > 3 and "---" in block:
                    add_table_from_markdown(doc, block)
                
                # Check if subsection heading
                elif block.startswith("### "):
                    subsec_title = block.replace("### ", "").strip()
                    create_heading(doc, subsec_title, level=2)
                
                # Regular paragraph
                else:
                    add_body_paragraph(doc, block)
    
    # References section
    doc.add_page_break()
    create_heading(doc, "VI. REFERENCES", level=1)
    
    references = [
        "[1] Verizon Media (2023). Data Breach Investigations Report. Available: verizon.com/dbir",
        "[2] Monrose, F., & Rubin, A. D. (2000). Keystroke dynamics as a biometric for authentication. In Future Generation Computer Systems, 16(4), 351-359.",
        "[3] Killourhy, K. S., & Maxion, R. A. (2009). The computer-science keystroke-dynamics benchmark database. In Proc. IEEE Int. Carnahan Conf. Security Tech. (ICCST), Cambridge, MA, USA.",
        "[4] Teh, P. S., Teoh, A. B., & Ong, T. S. (2013). A survey of keystroke dynamics biometrics. The Open Signal Processing Journal, 6(1), 1-15.",
        "[7] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.",
    ]
    
    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.paragraph_format.first_line_indent = Inches(-0.25)
        ref_para.paragraph_format.left_indent = Inches(0.25)
        ref_para.paragraph_format.space_after = Pt(3)
        
        for run in ref_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
    
    # Save document
    doc.save(output_path)
    print(f"✓ Generated conference format DOCX: {output_path}")


if __name__ == "__main__":
    markdown_file = Path("docs/RESEARCH_PAPER_DRAFT.md")
    output_file = Path("docs/RESEARCH_PAPER_CONFERENCE_FORMAT.docx")
    
    if not markdown_file.exists():
        print(f"✗ Markdown file not found: {markdown_file}")
        exit(1)
    
    create_conference_docx(str(markdown_file), str(output_file))
    print("✓ Conference format paper generated successfully!")
